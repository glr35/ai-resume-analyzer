import re
import io
import os
import fitz  # PyMuPDF
import streamlit as st
from datetime import date
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

# ---------- UYGULAMA AYARLARI ----------
st.set_page_config(page_title="AI Özgeçmiş Analizcisi (Offline)", page_icon="🧠", layout="wide")
load_dotenv()

# Kişisel varsayılanlar (.env ile özelleştirilebilir)
STUDENT_NAME = os.getenv("STUDENT_NAME", "Güler Göçmen")
ROLE_TARGET  = os.getenv("ROLE_TARGET", "AI & Python Developer")
UNIVERSITY   = os.getenv("UNIVERSITY", "Manisa Celal Bayar Üniversitesi")
PROGRAM      = os.getenv("PROGRAM", "Computer Engineering")
INTERNSHIP_START = os.getenv("INTERNSHIP_START", "2026")

# ---------- PDF'TEN METİN ÇIKARIMI ----------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    chunks = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for p in doc:
            chunks.append(p.get_text("text"))
    return "\n".join(chunks).strip()

# ---------- BASİT YARDIMCILAR ----------
LANG_PAT = re.compile(r"\b(english|turkish|german|french|italian|albanian|kurdish)\b", re.I)
TECH_KEYWORDS = {
    "python": ["python"],
    "java": ["java"],
    "c++": ["c\\+\\+"],
    "c": [r"\bc\b"],
    "js": ["javascript", r"\bjs\b"],
    "react": ["react", "react native"],
    "node": ["node"],
    "sql": ["sql", "postgres", "mysql"],
    "ml": ["machine learning", r"\bml\b", "scikit", "sklearn", "numpy", "pandas"],
    "dl": ["deep learning", "pytorch", "tensorflow", "keras"],
    "cv": ["computer vision", "opencv", "image", "yolo"],
    "nlp": ["nlp", "transformer", "bert", "hugging face", "spacy"],
    "api": ["api", "rest", "fastapi", "flask", "django"],
    "cloud": ["aws", "gcp", "azure"],
    "git": ["git", "github", "gitlab"],
}

def find_hits(text: str, patterns):
    found = set()
    for k, pats in patterns.items():
        for p in pats:
            if re.search(p, text, re.I):
                found.add(k)
                break
    return sorted(found)

def guess_languages(text: str):
    return sorted(set(m.group(0).capitalize() for m in LANG_PAT.finditer(text)))

def bullet(items):
    return "\n".join([f"- {i}" for i in items])

# ---------- ANALİZ MOTORU (OFFLINE) ----------
def analyze_cv(text: str):
    text_low = text.lower()

    tech = find_hits(text_low, TECH_KEYWORDS)
    langs = guess_languages(text_low)
    exp_years = 0
    for y in re.findall(r"(20\d{2}|19\d{2})", text):
        try:
            y = int(y)
            if 2000 <= y <= date.today().year:
                exp_years = max(exp_years, date.today().year - y)
        except:
            pass

    strengths = []
    if "python" in tech: strengths.append("Python ile uygulama geliştirme")
    if "ml" in tech or "dl" in tech: strengths.append("Makine öğrenmesi için NumPy/Pandas/Scikit deneyimi")
    if "cv" in tech: strengths.append("Görüntü işleme / OpenCV ile deneyim")
    if "nlp" in tech: strengths.append("Doğal dil işleme (Transformer/BERT) bilgisi")
    if "react" in tech: strengths.append("React/React Native ile ön-yüz/mobil geliştirme")
    if "api" in tech: strengths.append("REST API (Flask/FastAPI) ile servis geliştirme")
    if "git" in tech: strengths.append("Git ve GitHub akışlarına hakimiyet")
    if not strengths:
        strengths.append("Temel programlama ve algoritma bilgisi")
    strengths.append("Araştırmacı ve proje odaklı çalışma yaklaşımı")

    improve = []
    if "cloud" not in tech: improve.append("Bulut (AWS/GCP) üzerinde basit bir deploy (FastAPI + Docker)")
    if "sql" not in tech: improve.append("SQL ve temel veri modelleri")
    if "cv" not in tech and "nlp" not in tech and "ml" not in tech:
        improve.append("Bir yapay zeka alanında mini proje (CV veya NLP)")
    improve.append("Unit test ve basit CI (GitHub Actions)")
    if "react" not in tech: improve.append("Basit bir React/Next.js portföy")

    roles = [
        "AI/ML Intern (Computer Vision veya NLP)",
        "Python Backend Intern (FastAPI/Django)",
        "Data Intern (Pandas, SQL, raporlama)",
        "Mobile Intern (React Native)",
        "Automation/Script Intern (Python + GitHub Actions)",
    ]

    projects = []
    if "cv" in tech:
        projects.append("Realtime Image Captioner: OpenCV + küçük bir CNN/Transformer, Streamlit arayüzü")
    if "nlp" in tech:
        projects.append("Türkçe Duygu Analizi: küçük veri + Logistic Regression/Transformer, FastAPI servis")
    projects.append("CV Analiz Aracı: PDF’ten bilgi çıkaran ve öneri üreten Streamlit uygulaması (bu proje)")
    if "react" in tech:
        projects.append("React Native ile AI destekli metin özetleyici mobil uygulama")

    langs = langs or ["Turkish"]
    letter = (
        f"Dear Hiring Team,\n\n"
        f"I am {STUDENT_NAME}, a final-year {PROGRAM} student at {UNIVERSITY}. "
        f"My focus is {ROLE_TARGET}. I am seeking a long-term internship starting {INTERNSHIP_START}. "
        f"I have hands-on experience with {', '.join(tech) if tech else 'Python and core CS concepts'}, "
        f"and I enjoy building real products. I will be happy to contribute, learn fast and deliver.\n\n"
        f"Best regards,\n{STUDENT_NAME}"
    )

    return {
        "tech": tech,
        "languages": langs,
        "experience_years_est": exp_years,
        "strengths": strengths,
        "improve": improve,
        "roles": roles,
        "projects": projects,
        "letter": letter,
    }

# ---------- DOCX OLUŞTUR ----------
def build_docx(analysis, raw_text):
    doc = Document()
    styles = doc.styles['Normal']
    styles.font.name = 'Calibri'
    styles.font.size = Pt(11)

    doc.add_heading(f"{STUDENT_NAME} — CV Review", level=1)
    doc.add_paragraph(f"Target Role: {ROLE_TARGET}")
    doc.add_paragraph(f"University: {UNIVERSITY} | Program: {PROGRAM} | Internship: {INTERNSHIP_START}")
    doc.add_paragraph("")

    doc.add_heading("Key Strengths", level=2)
    for s in analysis["strengths"]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Areas to Improve", level=2)
    for s in analysis["improve"]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Suggested Roles", level=2)
    for s in analysis["roles"]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Mini Project Ideas", level=2)
    for s in analysis["projects"]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Motivation Letter (EN)", level=2)
    doc.add_paragraph(analysis["letter"])

    doc.add_heading("Extracted CV Text (raw)", level=2)
    doc.add_paragraph(raw_text[:5000])  # çok uzunsa kırp

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ---------- UI ----------
st.title("🧠 AI Özgeçmiş Analizcisi — v1.7 (Offline)")
st.caption("Tamamen offline çalışır. PDF'ten metin çıkarır, kural-tabanlı analiz üretir, DOCX indirir.")

tab1, tab2 = st.tabs(["📄 CV Analizi", "💾 Çıktı"])

with tab1:
    up = st.file_uploader("PDF CV yükle", type=["pdf"])
    cv_text = None
    if up:
        data = up.read()
        cv_text = extract_text_from_pdf(data)
        st.text_area("Çıkarılan CV Metni", cv_text, height=300)

        if st.button("⚙️ Analiz Et (Offline)"):
            if not cv_text or len(cv_text) < 20:
                st.error("CV metni çok kısa görünüyor. PDF'i kontrol et.")
            else:
                result = analyze_cv(cv_text)
                st.success("Analiz hazır ✅")
                st.subheader("✅ Güçlü Yönler")
                st.markdown(bullet(result["strengths"]))
                st.subheader("🛠️ Geliştirilmesi Gerekenler")
                st.markdown(bullet(result["improve"]))
                st.subheader("💼 Uygun Roller")
                st.markdown(bullet(result["roles"]))
                st.subheader("💡 Mini Proje Fikirleri")
                st.markdown(bullet(result["projects"]))
                st.subheader("✉️ Motivasyon Mektubu (EN)")
                st.code(result["letter"])

                st.session_state["analysis"] = result
                st.session_state["raw_text"] = cv_text

with tab2:
    st.subheader("DOCX Çıktısı")
    if "analysis" in st.session_state:
        buf = build_docx(st.session_state["analysis"], st.session_state["raw_text"])
        st.download_button(
            "⬇️ DOCX indir",
            data=buf,
            file_name=f"{STUDENT_NAME.lower().replace(' ', '_')}_cv_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.info("Önce 'Analiz Et (Offline)' ile analiz üret.")
